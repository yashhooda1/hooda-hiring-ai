"""
resume_builder.py — Deterministic document assembly for HoodaAgents.

No LLM calls here. Pure, testable functions:
  - detect_font(path)          : best-effort font family from PDF or DOCX
  - build_resume_docx(...)     : clean single-column ATS resume in a given font
  - build_cover_letter_docx(...)
  - tailor_existing_docx(...)  : font-PRESERVING surgical edit (DOCX upload only)

Font-preservation reality:
  * DOCX upload -> font is recoverable and preserved EXACTLY (run surgery).
  * PDF upload  -> fonts are often CID-subset with opaque names (e.g. "CIDFont+F4"),
                   so the family is unrecoverable. We fall back to a clean ATS font.
"""

import re
import copy
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ATS_FALLBACK_FONT = "Calibri"
# PDF generic/CID names we should NOT trust as a real family
_UNRESOLVABLE = re.compile(r"^(CIDFont|[A-Z0-9]+\+?F\d+|Font\d*)$", re.IGNORECASE)


def _normalize_font(raw):
    if not raw:
        return None
    name = re.sub(r"^[A-Z]{6}\+", "", raw)          # strip subset prefix ABCDEE+
    name = re.split(r"[-,]", name)[0]               # drop "-BoldMT" etc.
    name = re.sub(r"(MT|PS|Regular|Bold|Italic|Light|Medium)$", "", name)
    name = name.strip()
    if not name or _UNRESOLVABLE.match(name):
        return None
    return name


def detect_font(path):
    """Return a usable font family, or ATS_FALLBACK_FONT if unrecoverable."""
    p = path.lower()
    try:
        if p.endswith(".docx"):
            doc = Document(path)
            f = doc.styles["Normal"].font.name
            return _normalize_font(f) or ATS_FALLBACK_FONT
        if p.endswith(".pdf"):
            import pdfplumber
            from collections import Counter
            counter = Counter()
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    for ch in page.chars:
                        counter[ch.get("fontname", "")] += 1
            for raw, _ in counter.most_common():
                fam = _normalize_font(raw)
                if fam:
                    return fam
    except Exception:
        pass
    return ATS_FALLBACK_FONT


def _apply_base_font(doc, font_name):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(10.5)


def _heading(doc, text, font_name):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.name = font_name
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    # thin rule under heading
    p_fmt = p.paragraph_format
    p_fmt.space_after = Pt(2)
    return p


def _skills_lines(skills):
    """Yield (label, text). dict -> grouped lines; list/str -> single line (label None)."""
    if isinstance(skills, dict):
        for cat, vals in skills.items():
            if vals:
                yield cat, (", ".join(vals) if isinstance(vals, list) else str(vals))
    elif isinstance(skills, list):
        if skills:
            yield None, ", ".join(skills)
    elif skills:
        yield None, str(skills)


def build_resume_docx(data, font_name, out=None):
    """data: enriched structured dict (see schema in resume_generator). Bytes if out is None."""
    doc = Document()
    _apply_base_font(doc, font_name)
    doc.styles["Normal"].font.size = Pt(9.5)  # compact, aims for one page
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(30)
        section.left_margin = section.right_margin = Pt(44)

    def center(text, size, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.name = font_name; r.font.size = Pt(size)

    center(data.get("name", ""), 18, bold=True)
    if data.get("tagline"):
        center(data["tagline"], 10.5, bold=True)
    if data.get("contact"):
        center(data["contact"], 9)
    if data.get("links"):
        center(data["links"], 9)
    if data.get("work_authorization"):
        center(data["work_authorization"], 8.5, italic=True)

    if data.get("summary"):
        _heading(doc, "Summary", font_name)
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        _heading(doc, "Technical Skills", font_name)
        for label, txt in _skills_lines(data["skills"]):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            if label:
                p.add_run(f"{label}  ").bold = True
            p.add_run(txt)

    if data.get("experience"):
        _heading(doc, "Experience", font_name)
        for job in data["experience"]:
            h = doc.add_paragraph()
            h.paragraph_format.space_after = Pt(0)
            h.add_run(f"{job.get('title','')} — {job.get('company','')}").bold = True
            meta = " | ".join(x for x in [job.get("location", ""), job.get("dates", "")] if x)
            if meta:
                h.add_run(f"   {meta}").italic = True
            for b in job.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    if data.get("projects"):
        _heading(doc, "Selected Projects", font_name)
        for pr in data["projects"]:
            h = doc.add_paragraph()
            h.paragraph_format.space_after = Pt(0)
            title = pr.get("name", "")
            if pr.get("subtitle"):
                title += f" — {pr['subtitle']}"
            h.add_run(title).bold = True
            if pr.get("tech"):
                h.add_run(f"   {pr['tech']}").italic = True
            for b in pr.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    if data.get("education") or data.get("certifications"):
        _heading(doc, "Education & Certifications", font_name)
        for ed in data.get("education", []):
            e = doc.add_paragraph()
            e.paragraph_format.space_after = Pt(1)
            e.add_run(ed.get("degree", "")).bold = True
            tail = ", ".join(x for x in [ed.get("school", ""), ed.get("location", ""), ed.get("dates", "")] if x)
            if tail:
                e.add_run(f" — {tail}")
        certs = data.get("certifications")
        if certs:
            text = "  ·  ".join(certs) if isinstance(certs, list) else str(certs)
            doc.add_paragraph(text)

    return _save(doc, out)


# ---- PDF builders (fpdf2, no system deps) ----
_SERIF_HINTS = ("times", "georgia", "garamond", "cambria", "serif", "tex",
                "computer modern", "cmr", "minion", "palatino", "book antiqua")
_PUNCT = {"\u2014": "-", "\u2013": "-", "\u2022": "-", "\u25cf": "-", "\u2011": "-",
          "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
          "\u2026": "...", "\u00a0": " "}


def _pdf_family(font_name):
    n = (font_name or "").lower()
    return "Times" if any(h in n for h in _SERIF_HINTS) else "Helvetica"


def _pdf_safe(s):
    """Core PDF fonts are latin-1. Convert typographic punctuation to ASCII,
    keep Western accents (é, ñ, ü...), drop anything else rather than crash."""
    s = s or ""
    for k, v in _PUNCT.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "ignore").decode("latin-1")


def _pdf_out(pdf, out):
    data = bytes(pdf.output())
    if out is None:
        return data
    with open(out, "wb") as f:
        f.write(data)
    return out


def build_resume_pdf(data, font_name, out=None):
    from fpdf import FPDF
    fam = _pdf_family(font_name)
    pdf = FPDF(format="Letter")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_margins(16, 12, 16)
    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin

    def center(text, size, style=""):
        pdf.set_font(fam, style, size)
        pdf.multi_cell(W, size * 0.42 + 1, _pdf_safe(text), align="C")

    center(data.get("name", ""), 18, "B")
    if data.get("tagline"):
        center(data["tagline"], 10.5, "B")
    if data.get("contact"):
        center(data["contact"], 8.5)
    if data.get("links"):
        center(data["links"], 8.5)
    if data.get("work_authorization"):
        center(data["work_authorization"], 8, "I")

    def heading(t):
        pdf.ln(1.8)
        pdf.set_font(fam, "B", 11)
        pdf.multi_cell(W, 5.5, _pdf_safe(t.upper()))
        y = pdf.get_y()
        pdf.set_draw_color(170, 170, 170)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(0.8)

    def body(t, size=9.5):
        pdf.set_font(fam, "", size)
        pdf.multi_cell(W, 4.4, _pdf_safe(t))

    def bullet(t):
        pdf.set_font(fam, "", 9.5)
        pdf.multi_cell(W, 4.4, _pdf_safe("- " + t))

    if data.get("summary"):
        heading("Summary"); body(data["summary"])

    if data.get("skills"):
        heading("Technical Skills")
        for label, txt in _skills_lines(data["skills"]):
            pdf.set_font(fam, "B", 9.5)
            if label:
                pdf.write(4.4, _pdf_safe(f"{label}  "))
            pdf.set_font(fam, "", 9.5)
            pdf.write(4.4, _pdf_safe(txt))
            pdf.ln(4.4)

    if data.get("experience"):
        heading("Experience")
        for job in data["experience"]:
            pdf.set_font(fam, "B", 9.5)
            line = f"{job.get('title','')} - {job.get('company','')}"
            meta = " | ".join(x for x in [job.get("location", ""), job.get("dates", "")] if x)
            pdf.multi_cell(W, 4.6, _pdf_safe(line + (f"   {meta}" if meta else "")))
            for b in job.get("bullets", []):
                bullet(b)
            pdf.ln(0.8)

    if data.get("projects"):
        heading("Selected Projects")
        for pr in data["projects"]:
            title = pr.get("name", "")
            if pr.get("subtitle"):
                title += f" - {pr['subtitle']}"
            if pr.get("tech"):
                title += f"   {pr['tech']}"
            pdf.set_font(fam, "B", 9.5)
            pdf.multi_cell(W, 4.6, _pdf_safe(title))
            for b in pr.get("bullets", []):
                bullet(b)
            pdf.ln(0.8)

    if data.get("education") or data.get("certifications"):
        heading("Education & Certifications")
        for ed in data.get("education", []):
            pdf.set_font(fam, "B", 9.5)
            tail = ", ".join(x for x in [ed.get("school", ""), ed.get("location", ""), ed.get("dates", "")] if x)
            pdf.multi_cell(W, 4.6, _pdf_safe(ed.get("degree", "") + (f" - {tail}" if tail else "")))
        certs = data.get("certifications")
        if certs:
            body("  ·  ".join(certs) if isinstance(certs, list) else str(certs))

    return _pdf_out(pdf, out)


def build_cover_letter_pdf(body_text, font_name, out=None):
    from fpdf import FPDF
    fam = _pdf_family(font_name)
    pdf = FPDF(format="Letter")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(25, 20, 25)
    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font(fam, "", 11)
    for para in body_text.split("\n\n"):
        para = para.strip()
        if para:
            pdf.multi_cell(W, 6, _pdf_safe(para))
            pdf.ln(3)
    return _pdf_out(pdf, out)


# ---- Faithful DOCX -> PDF via LibreOffice (exact render; needs soffice installed) ----
import os
import shutil
import subprocess
import tempfile


def libreoffice_available():
    return shutil.which("soffice") is not None or shutil.which("libreoffice") is not None


def docx_to_pdf(docx_bytes, out=None):
    """Convert DOCX bytes to PDF using headless LibreOffice.

    Produces a PDF that matches the DOCX exactly (layout + font, provided the
    font — or a metric-compatible substitute — is installed in the container).
    Raises RuntimeError if soffice is unavailable or conversion fails; callers
    should fall back to the fpdf2 builder.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) is not installed")

    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.docx")
        with open(src, "wb") as f:
            f.write(docx_bytes)
        # Unique profile dir avoids "another instance is running" locks under concurrency
        profile = os.path.join(tmp, "profile")
        cmd = [
            soffice, "--headless", "--nologo", "--nofirststartwizard",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to", "pdf", "--outdir", tmp, src,
        ]
        proc = subprocess.run(cmd, capture_output=True, timeout=90)
        pdf_path = os.path.join(tmp, "in.pdf")
        if proc.returncode != 0 or not os.path.exists(pdf_path):
            raise RuntimeError(
                f"LibreOffice conversion failed: {proc.stderr.decode('utf-8', 'ignore')[:300]}"
            )
        with open(pdf_path, "rb") as f:
            data = f.read()

    if out is None:
        return data
    with open(out, "wb") as f:
        f.write(data)
    return out


def build_cover_letter_docx(body_text, font_name, out=None):
    doc = Document()
    _apply_base_font(doc, font_name)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(54)
        section.left_margin = section.right_margin = Pt(72)
    for para in body_text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    return _save(doc, out)


# ---- Font-PRESERVING surgical path (DOCX upload only) ----
def _replace_paragraph_text(paragraph, new_text):
    runs = paragraph.runs
    rPr = None
    if runs:
        src = runs[0]._element.rPr
        if src is not None:
            rPr = copy.deepcopy(src)
        for r in list(runs):
            r._element.getparent().remove(r._element)
    nr = paragraph.add_run(new_text)
    if rPr is not None:
        nr._element.insert(0, rPr)


def _iter_paragraphs(doc):
    from docx.document import Document as _Doc
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    def walk(parent):
        elm = parent.element.body if isinstance(parent, _Doc) else parent._element
        for child in elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                for row in Table(child, parent).rows:
                    for cell in row.cells:
                        yield from walk(cell)
    yield from walk(doc)


def extract_docx_paragraphs(path):
    doc = Document(path)
    return [{"index": i, "text": p.text} for i, p in enumerate(_iter_paragraphs(doc)) if p.text.strip()]


def tailor_existing_docx(path, edits, out=None):
    """edits: [{'index', 'new_text'}] -> font preserved exactly."""
    doc = Document(path)
    by_index = {e["index"]: e["new_text"] for e in edits}
    for i, p in enumerate(_iter_paragraphs(doc)):
        if i in by_index:
            _replace_paragraph_text(p, by_index[i])
    return _save(doc, out)


def extract_text(path):
    """Unified text extraction for the compare step. Handles PDF and DOCX."""
    p = path.lower()
    if p.endswith(".docx"):
        doc = Document(path)
        return "\n".join(para.text for para in _iter_paragraphs(doc))
    # default: PDF via pdfplumber
    import pdfplumber
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text


def to_markdown(data):
    """Render the tailored resume dict as Markdown / plain text for easy copy-paste
    (e.g. back into a LaTeX master or another editor)."""
    lines = []
    if data.get("name"):
        lines.append(f"# {data['name']}")
    for key in ("tagline", "contact", "links", "work_authorization"):
        if data.get(key):
            lines.append(data[key])

    if data.get("summary"):
        lines += ["", "## Summary", data["summary"]]

    if data.get("skills"):
        lines += ["", "## Technical Skills"]
        for label, txt in _skills_lines(data["skills"]):
            lines.append(f"**{label}:** {txt}" if label else txt)

    if data.get("experience"):
        lines += ["", "## Experience"]
        for job in data["experience"]:
            meta = " | ".join(x for x in [job.get("location", ""), job.get("dates", "")] if x)
            head = f"**{job.get('title','')} — {job.get('company','')}**"
            if meta:
                head += f" — {meta}"
            lines += ["", head]
            for b in job.get("bullets", []):
                lines.append(f"- {b}")

    if data.get("projects"):
        lines += ["", "## Selected Projects"]
        for pr in data["projects"]:
            title = pr.get("name", "")
            if pr.get("subtitle"):
                title += f" — {pr['subtitle']}"
            head = f"**{title}**"
            if pr.get("tech"):
                head += f" — {pr['tech']}"
            lines += ["", head]
            for b in pr.get("bullets", []):
                lines.append(f"- {b}")

    if data.get("education") or data.get("certifications"):
        lines += ["", "## Education & Certifications"]
        for ed in data.get("education", []):
            tail = ", ".join(x for x in [ed.get("school", ""), ed.get("location", ""), ed.get("dates", "")] if x)
            lines.append(f"**{ed.get('degree','')}**" + (f" — {tail}" if tail else ""))
        for c in (data.get("certifications") or []):
            lines.append(f"- {c}")

    return "\n".join(lines).strip() + "\n"


def _save(doc, out):
    if out is None:
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    doc.save(out)
    return out
